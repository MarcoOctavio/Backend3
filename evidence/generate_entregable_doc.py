from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Entregable_Adoptme_Backend.docx"


def read(rel):
    return (ROOT / rel).read_text(encoding="utf-8", errors="replace")


def generated_tree():
    excluded_dirs = {".git", "node_modules", "evidence"}
    excluded_files = {".DS_Store"}
    lines = ["."]

    def visible_children(path):
        children = []
        for child in sorted(path.iterdir(), key=lambda p: (p.is_file(), p.name.lower())):
            if child.name in excluded_dirs or child.name in excluded_files:
                continue
            if child.is_file() and child.name in excluded_files:
                continue
            if str(child.relative_to(ROOT)).startswith("public/documents/") and child.name != ".gitkeep":
                continue
            children.append(child)
        return children

    def walk(path, prefix=""):
        children = visible_children(path)
        for index, child in enumerate(children):
            last = index == len(children) - 1
            connector = "└── " if last else "├── "
            rel_name = child.name + ("/" if child.is_dir() else "")
            lines.append(prefix + connector + rel_name)
            if child.is_dir():
                walk(child, prefix + ("    " if last else "│   "))

    walk(ROOT)
    return "\n".join(lines)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code_block(doc, text, max_font=8):
    for line in text.rstrip("\n").splitlines() or [""]:
        p = doc.add_paragraph(style="CodeBlock")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(max_font)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, RGBColor(0, 0, 0)),
        ("Heading 2", 16, 18, 6, RGBColor(0, 0, 0)),
        ("Heading 3", 14, 16, 4, RGBColor(67, 67, 67)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code.font.size = Pt(8)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.15)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Entregable del proyecto Adoptme Backend")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(26)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    subtitle = doc.add_paragraph("Documento tecnico para Google Docs")
    subtitle.runs[0].font.name = "Arial"
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    doc.add_paragraph("Proyecto: Adoptme Backend API")
    doc.add_paragraph("Imagen Docker: marcoven/adoptme-backend:2.0")


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        hdr[idx].text = value
        set_cell_shading(hdr[idx], "F8F9FA")
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    return table


def build():
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("1. Estructura del proyecto", level=1)
    doc.add_paragraph(
        "El repositorio implementa una API REST organizada por capas. Las rutas reciben las solicitudes HTTP, "
        "los controladores coordinan la respuesta, los servicios concentran la logica de negocio y los repositorios/DAO "
        "encapsulan el acceso a MongoDB mediante modelos Mongoose."
    )
    doc.add_heading("Arbol de directorios", level=2)
    doc.add_paragraph("Comando equivalente: find . -maxdepth 4, excluyendo .git, node_modules, evidence y archivos generados.")
    add_code_block(doc, generated_tree(), max_font=7)
    doc.add_heading("Proposito de archivos y carpetas principales", level=2)
    add_table(
        doc,
        [
            ("Elemento", "Proposito"),
            ("src/app.js", "Configura Express, Swagger, middlewares, conexion a MongoDB y routers."),
            ("src/routes", "Define los endpoints HTTP por dominio: users, pets, sessions, mocks, logger y adoptions."),
            ("src/controllers", "Contiene la capa request/response y delega la logica en servicios."),
            ("src/services", "Agrupa operaciones de negocio reutilizables."),
            ("src/repository", "Abstrae operaciones comunes de persistencia."),
            ("src/dao y src/dao/models", "Implementa acceso a datos y esquemas Mongoose."),
            ("src/docs", "Contiene especificaciones Swagger y ejemplos de cURL."),
            ("src/middlewares", "Incluye logger de requests y manejador centralizado de errores."),
            ("test", "Contiene tests funcionales con Mocha, Chai y Supertest."),
            ("Dockerfile", "Define la imagen productiva multi-stage para ejecutar la API."),
            ("README.md", "Documentacion reproducible del proyecto."),
        ],
    )

    doc.add_heading("2. Tests funcionales", level=1)
    doc.add_paragraph(
        "Los tests funcionales de adopciones cubren todos los endpoints definidos en src/routes/adoption.router.js: "
        "listar adopciones, consultar una adopcion por ID y crear una adopcion entre usuario y mascota."
    )
    doc.add_heading("Codigo completo de test/adoptions.test.js", level=2)
    add_code_block(doc, read("test/adoptions.test.js"), max_font=7)
    doc.add_heading("Explicacion de los grupos de tests", level=2)
    add_bullets(
        doc,
        [
            "Preparacion: crea un usuario mediante POST /api/sessions/register y una mascota mediante POST /api/pets para obtener IDs validos.",
            "Creacion de adopcion: valida que POST /api/adoptions/:uid/:pid responda 200, status success y mensaje Pet adopted.",
            "Listado de adopciones: valida GET /api/adoptions, confirma que el payload sea un array y localiza la adopcion creada.",
            "Consulta por ID: valida GET /api/adoptions/:aid y verifica que el payload corresponda al ID recuperado.",
            "Reglas de negocio negativas: valida que una mascota ya adoptada responda 400 y que un usuario inexistente responda 404.",
        ],
    )
    doc.add_heading("Evidencia de ejecucion de tests", level=2)
    add_code_block(doc, read("evidence/test-log.txt"), max_font=7)

    doc.add_heading("3. Dockerizacion", level=1)
    doc.add_heading("Dockerfile completo", level=2)
    add_code_block(doc, read("Dockerfile"), max_font=8)
    doc.add_heading("Decisiones de optimizacion", level=2)
    add_bullets(
        doc,
        [
            "Imagen base node:20-alpine: reduce el tamano final respecto a imagenes Debian completas y mantiene compatibilidad con Node 20.",
            "Multi-stage build: separa la instalacion de dependencias de la imagen runner final.",
            "npm ci --omit=dev: instala dependencias reproducibles desde package-lock.json y excluye librerias de testing/desarrollo.",
            "dotenv esta clasificado como dependencia de produccion porque src/app.js lo importa durante el arranque del servidor.",
            "Copia selectiva: la etapa final copia solo node_modules, package*.json y src, evitando archivos innecesarios.",
            "NODE_ENV=production: declara el modo productivo del runtime.",
            "USER node: ejecuta el proceso con usuario no root para reducir privilegios dentro del contenedor.",
            ".dockerignore: evita enviar node_modules, .git y logs al contexto de build.",
        ],
    )
    doc.add_heading("Log de construccion de imagen Docker", level=2)
    add_code_block(doc, read("evidence/docker-build-log.txt"), max_font=6)

    doc.add_heading("4. Imagen Docker", level=1)
    doc.add_paragraph("Nombre y tag de la imagen generada: marcoven/adoptme-backend:2.0")
    doc.add_heading("Evidencia de imagen construida", level=2)
    add_code_block(doc, read("evidence/docker-image-log.txt"), max_font=8)
    doc.add_heading("Evidencia de ejecucion del contenedor", level=2)
    add_code_block(doc, read("evidence/docker-container-status.txt"), max_font=8)
    add_code_block(doc, read("evidence/docker-run-log.txt"), max_font=8)

    doc.add_heading("5. Ejecucion del proyecto", level=1)
    doc.add_heading("Construir la imagen Docker", level=2)
    add_steps(doc, ["Ejecutar docker build -t marcoven/adoptme-backend:2.0 . desde la raiz del proyecto."])
    doc.add_heading("Ejecutar el contenedor", level=2)
    add_steps(
        doc,
        [
            "Verificar que MongoDB este disponible.",
            "Ejecutar docker run --name adoptme-backend -p 8080:8080 -e PORT=8080 -e MONGO_URI=mongodb://host.docker.internal:27017/adoptme -e JWT_SECRET=superSecretKey marcoven/adoptme-backend:2.0.",
            "Abrir http://localhost:8080/api/docs/ para verificar Swagger UI.",
        ],
    )
    doc.add_heading("Correr los tests", level=2)
    add_steps(doc, ["Instalar dependencias con npm install.", "Configurar .env.", "Ejecutar npm test."])
    doc.add_heading("Evidencia de ejecucion exitosa", level=2)
    add_code_block(doc, read("evidence/project-http-log.txt"), max_font=6)

    doc.add_heading("6. README", level=1)
    doc.add_paragraph("Contenido completo del README.md actualizado:")
    add_code_block(doc, read("README.md"), max_font=7)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
